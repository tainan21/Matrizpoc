from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
ASSETS = ROOT / "docs" / "seumei" / "assets"
OUTPUT = ROOT / "docs" / "seumei" / "Seumei-Progress-and-Roadmap-2026-08-24.docx"

PURPLE = "8B5CF6"
PURPLE_DARK = "5B25C5"
INK = "121526"
MUTED = "5E6478"
LIGHT = "F4F2FB"
MID = "E8E3F5"
GREEN = "177245"
AMBER = "9A6700"
RED = "A12828"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_font(run, size=9, color=MUTED)
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr_text, fld_end])


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_body(doc, text, bold_lead=None):
    paragraph = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        first = paragraph.add_run(bold_lead)
        set_font(first, bold=True)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_font(rest)
    else:
        run = paragraph.add_run(text)
        set_font(run)
    return paragraph


def add_status_callout(doc, label, text, color=PURPLE, fill=LIGHT):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    label_run = paragraph.add_run(f"{label}  ")
    set_font(label_run, size=10.5, color=color, bold=True)
    text_run = paragraph.add_run(text)
    set_font(text_run, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.167
        run = p.add_run(item)
        set_font(run)


def add_table(doc, headers, rows, widths_dxa, status_col=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_tr_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_tr_pr.append(repeat_header)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, MID)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_font(run, size=9.5, color=INK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            status_color = None
            if status_col == index:
                lower = str(value).lower()
                status_color = GREEN if any(k in lower for k in ("conclu", "implement", "preserv", "passou", "ativo")) else AMBER
            set_font(run, size=9.2, color=status_color or INK, bold=status_col == index)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_figure(doc, filename, caption, width, alt_text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    shape = run.add_picture(str(ASSETS / filename), width=Inches(width))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt_text)
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_before = Pt(4)
    caption_p.paragraph_format.space_after = Pt(10)
    caption_p.paragraph_format.keep_with_next = False
    caption_run = caption_p.add_run(caption)
    set_font(caption_run, size=9, color=MUTED, italic=True)


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, PURPLE_DARK, 16, 8),
        ("Heading 2", 13, PURPLE_DARK, 12, 6),
        ("Heading 3", 12, PURPLE_DARK, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = doc.styles[list_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def configure_running_furniture(doc):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("SEUMEI BUSINESS OS  /  PROGRESSO E ROADMAP")
    set_font(run, size=8.5, color=MUTED, bold=True)
    footer = section.footer
    add_page_number(footer.paragraphs[0])


def build():
    doc = Document()
    configure_styles(doc)
    configure_running_furniture(doc)
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    cover = doc.add_paragraph()
    cover.paragraph_format.space_before = Pt(44)
    cover.paragraph_format.space_after = Pt(14)
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker = cover.add_run("RELATÓRIO DE IMPLEMENTAÇÃO  ·  24 AGO 2026")
    set_font(kicker, size=10, color=PURPLE, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    title_run = title.add_run("Seumei Business OS")
    set_font(title_run, size=30, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(20)
    subtitle_run = subtitle.add_run("Fundação multiempresa, shell inteligente e ponto de retomada")
    set_font(subtitle_run, size=15, color=MUTED)

    add_figure(
        doc,
        "seumei-hub-desktop.png",
        "Seumei Hub em execução com Galáxia Burger e Matriz Labs isoladas.",
        6.5,
        "Captura do Seumei Hub em desktop com duas empresas e aplicativos disponíveis.",
    )
    cover_meta = doc.add_paragraph()
    cover_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_meta.paragraph_format.space_before = Pt(6)
    meta_run = cover_meta.add_run("Status: fundação aprovada, integrada e publicada na main (90a377d)")
    set_font(meta_run, size=10.5, color=GREEN, bold=True)
    doc.add_page_break()

    add_heading(doc, "1. Resumo executivo")
    add_status_callout(
        doc,
        "MARCO ATUAL",
        "A Seumei deixou de depender de uma seleção visual de empresa e passou a resolver operações por usuário autenticado, membership validada e tenant confiável.",
    )
    add_body(doc, "Esta entrega consolidou a fundação pós-login da Seumei sem reconstruir o produto do zero. A arquitetura existente que já respeitava o monorepo Matriz foi preservada; os pontos frágeis de estado global e contexto único foram substituídos por domínios locais, contratos explícitos e repositórios vinculados ao tenant.")
    add_body(doc, "O modo demo agora é a porta de entrada padrão para testes práticos. A conta demo navega pela mesma cadeia de autenticação e autorização da aplicação, participa de Galáxia Burger e Matriz Labs e mantém dados operacionais e aplicativos separados entre as empresas.")
    add_bullets(doc, [
        "Merge local concluído sem conflitos e publicado na branch main.",
        "Conta demo canônica: demo@seumei.local, com acesso de um clique.",
        "Hub, troca de empresa, registro de aplicativos e shell compartilhado implementados.",
        "Preferência pessoal de aparência separada de branding da empresa e apresentação da loja.",
        "20/20 testes Seumei e 148/148 verificações globais aprovados na árvore publicada.",
    ])

    add_heading(doc, "2. Arquitetura encontrada e decisões de preservação")
    add_body(doc, "O repositório é um monorepo pnpm/Turborepo orientado a aplicativos. As leis arquiteturais exigem domínio forte dentro do aplicativo, comunicação entre apps apenas por contratos públicos e extração tardia para packages compartilhados. A refatoração respeitou esse modelo: o domínio Seumei permaneceu em apps/seumei e MatrizLib foi consumida como infraestrutura visual compartilhada.")
    add_table(
        doc,
        ["Área descoberta", "Decisão", "Resultado"],
        [
            ("Autenticação Matriz", "Preservada", "O modo demo usa o broker e a sessão existentes."),
            ("Login Seumei", "Preservado", "Somente o acesso demo foi acrescentado; a experiência aceita não foi redesenhada."),
            ("Monorepo e manifests", "Preservados", "Nenhum import interno entre aplicativos foi introduzido."),
            ("MatrizLib", "Integrada", "ThemeToggle e superfícies de autenticação continuam compartilhados."),
            ("Estado pós-login", "Refatorado", "Tenant, membership e apps agora possuem contratos explícitos."),
            ("Navegação", "Consolidada", "Topbar, sidebar e app switcher pertencem ao shell Seumei."),
        ],
        [2300, 1600, 5460],
        status_col=1,
    )

    add_heading(doc, "3. Mapa de domínios")
    add_body(doc, "O mapa abaixo registra ownership e maturidade. Os limites seguem a realidade atual do aplicativo, sem criar um pacote de domínio universal nem antecipar complexidade de ERP.")
    add_table(
        doc,
        ["Domínio", "Ownership / estado atual", "Próximo uso"],
        [
            ("Identity / Access", "Autenticação compartilhada preservada", "Sessões reais e recuperação de conta"),
            ("Companies / Tenant", "Company e CompanyBranding implementados", "Onboarding e persistência"),
            ("Membership / Authorization", "Membership, TenantContext e policy implementados", "Permissões por operação"),
            ("Application Catalog", "AppDefinition tipado implementado", "Contribuições de comandos e rotas"),
            ("Installed Applications", "Instalação por company implementada", "Ativar/desativar apps"),
            ("Catalog / Products", "Limite definido; slice ainda pendente", "Produtos, categorias e modificadores"),
            ("Customers / CRM", "Limite definido; slice ainda pendente", "Customer List e Customer 360"),
            ("Orders / Commerce", "Limite definido; slice ainda pendente", "Carrinho, checkout e pedido"),
            ("Store", "Separação conceitual definida", "Resolver público, catálogo publicado"),
            ("Inventory / Finance", "Apps registrados; domínio pendente", "Sinais operacionais mínimos"),
            ("Marketing / Reporting", "Apps registrados; domínio pendente", "Leitura de dados compartilhados"),
            ("Platform / Shell", "Shell inteligente implementado", "Contribuições contextuais"),
            ("Preferences / Appearance", "Preferência pessoal implementada", "Persistência por usuário"),
        ],
        [2050, 3810, 3500],
    )

    add_heading(doc, "4. Arquitetura de tenant e isolamento")
    add_status_callout(
        doc,
        "CADEIA DE CONFIANÇA",
        "Usuário autenticado → membership validada → TenantContext resolvido → policy de aplicativo → repositório vinculado ao tenant.",
        color=GREEN,
        fill="EAF7F0",
    )
    add_body(doc, "O companyId recebido por uma rota ou componente não é tratado como autorização. Ele é apenas uma intenção de navegação, posteriormente confrontada com memberships válidas da sessão. Serviços e repositórios recebem contexto resolvido ou já nascem vinculados ao tenant.")
    add_bullets(doc, [
        "Galáxia Burger possui sete aplicativos instalados; Matriz Labs possui quatro.",
        "Pedidos, Estoque, Financeiro e Loja não aparecem nem abrem na Matriz Labs.",
        "Trocar de empresa atualiza o contexto operacional e o conjunto de navegação.",
        "A preferência de aparência continua pertencendo ao usuário, independentemente da empresa ativa.",
        "Repositórios globais de produtos, clientes e pedidos não serão permitidos nos próximos slices.",
    ])
    add_figure(
        doc,
        "seumei-tenant-access-denied.png",
        "Acesso direto a Pedidos na Matriz Labs falha fechado porque o aplicativo não está instalado/autorizado.",
        6.5,
        "Captura do estado de acesso indisponível ao tentar abrir Pedidos na Matriz Labs.",
    )

    add_heading(doc, "5. Modo demo como ambiente padrão")
    add_body(doc, "A conta demo foi criada para uso recorrente, testes funcionais e demonstrações sem misturar tenants. Ela entra pela autenticação existente e recebe memberships explícitas, em vez de contornar as regras com flags em componentes.")
    add_table(
        doc,
        ["Item", "Definição"],
        [
            ("Conta", "demo@seumei.local"),
            ("Entrada", "Botão Entrar no modo demo na experiência de login aceita"),
            ("Tenant primário", "Galáxia Burger · Proprietário · 7 apps"),
            ("Tenant secundário", "Matriz Labs · Administrador · 4 apps"),
            ("Preferência", "Tema do usuário; não é branding da empresa"),
            ("Dados", "Fixtures coerentes e independentes por tenant"),
        ],
        [2300, 7060],
    )

    add_heading(doc, "6. Shell autenticado e navegação")
    add_body(doc, "O shell agora é uma capacidade compartilhada do produto e não uma implementação duplicada por página. O app registry dirige atalhos e disponibilidade; as aplicações contribuem contexto sem recriar topbars.")
    add_bullets(doc, [
        "Topbar compacta com busca, apps, notificações, aparência e conta.",
        "Sidebar contextual com estado recolhido e expansão explícita/foco.",
        "App switcher persistente para alternar capacidades instaladas.",
        "Interações equivalentes para mouse, teclado, foco e touch.",
        "Transições curtas e remoção de movimento quando prefers-reduced-motion está ativo.",
    ])
    add_figure(
        doc,
        "seumei-galaxia-shell-desktop.png",
        "Shell operacional compartilhado da Galáxia Burger com sete capacidades instaladas.",
        6.5,
        "Captura do workspace da Galáxia Burger com topbar, sidebar compacta e cards de aplicativos.",
    )

    add_heading(doc, "7. Rotas entregues")
    add_table(
        doc,
        ["Rota", "Função", "Proteção"],
        [
            ("/login", "Login existente + acesso demo", "Autenticação Matriz"),
            ("/hub", "Empresas e capacidades disponíveis", "Usuário autenticado"),
            ("/c/[companySlug]", "Entrada operacional da empresa", "Membership válida"),
            ("/c/[companySlug]/apps/[appId]", "Entrada de aplicativo", "Membership + app policy"),
        ],
        [2800, 3300, 3260],
    )

    add_heading(doc, "8. Responsividade validada")
    add_body(doc, "A responsividade foi tratada como composição própria, não como simples redução do desktop. O Hub muda para coluna única, mantém as ações principais disponíveis e transforma navegação por proximidade em controles explícitos.")
    add_table(
        doc,
        ["Viewport", "Resultado"],
        [
            ("390 × 844", "Mobile: cards em coluna, controles touch e nenhum overflow horizontal."),
            ("768 × 1024", "Tablet: shell adaptativo e ações estáveis."),
            ("1180 × 820", "Desktop compacto: conteúdo íntegro e navegação acessível."),
            ("1440 × 900", "Desktop amplo: densidade e proporções próximas às referências."),
        ],
        [2100, 7260],
    )
    add_figure(
        doc,
        "seumei-hub-mobile.png",
        "Hub mobile em 390 × 844 com Galáxia Burger e Matriz Labs em composição vertical.",
        2.55,
        "Captura mobile longa do Seumei Hub com duas empresas e lista de aplicativos.",
    )

    add_heading(doc, "9. Verificação e evidências")
    add_table(
        doc,
        ["Verificação", "Resultado", "Escopo"],
        [
            ("Testes Seumei", "20/20 passaram", "TenantContext, policies, fixtures, presenter, demo e shell"),
            ("Typecheck", "Passou", "@matriz/app-seumei"),
            ("Lint", "Passou sem warnings", "@matriz/app-seumei"),
            ("Smoke global", "148/148 passaram", "Contratos, manifests, eventos, auth e boundaries"),
            ("Browser", "Passou", "Login demo, troca de empresa, rotas e responsividade"),
            ("Git", "Publicado", "main local/remota em 90a377d antes deste relatório"),
        ],
        [2200, 2150, 5010],
        status_col=1,
    )
    add_body(doc, "Nota de execução: uma tentativa de smoke no worktree temporário falhou porque junctions de node_modules não reproduzem todos os symlinks absolutos do pnpm. A árvore main foi comparada byte a byte com a árvore do worktree íntegro, onde as 148 verificações passaram. Não houve falha funcional no código publicado.")

    add_heading(doc, "10. Débito técnico e limites conhecidos")
    add_bullets(doc, [
        "Os repositórios atuais são fixtures em memória; a persistência real ainda deve adotar consultas obrigatoriamente vinculadas ao tenant.",
        "Products, Store, Orders, Dashboard e CRM ainda exibem entradas de capacidade, não slices operacionais completos.",
        "As permissões atuais cobrem acesso a apps; permissões finas por operação serão introduzidas com os casos de uso reais.",
        "A separação Web/Tauri está definida como contrato, mas os ports de filesystem, notificações e armazenamento seguro ainda não foram implementados.",
        "A fidelidade visual atingiu o Hub e o shell-base; as referências de Produtos, Store, Dashboard e CRM permanecem como metas dos próximos ciclos.",
    ])
    add_status_callout(
        doc,
        "RISCO PRINCIPAL",
        "Introduzir catálogo ou pedidos por arrays globais quebraria a garantia multi-tenant. O próximo slice deve começar por contratos, testes de isolamento e pricing de domínio antes da tela.",
        color=RED,
        fill="FCEEEE",
    )

    add_heading(doc, "11. Roadmap de implementação")
    add_table(
        doc,
        ["Slice", "Objetivo", "Definição de pronto"],
        [
            ("2 · Catalog / Products", "Product, Category, Modifier, disponibilidade e pricing", "Admin fiel à referência; dois tenants; edição altera a fonte única"),
            ("3 · Store / Commerce", "Resolver público, Product Detail, Cart e Order", "Preço fora do React; produto indisponível some da loja; pedido real criado"),
            ("4 · Orders / Dashboard / CRM", "Consumir o mesmo pedido e cliente", "Pedido aparece nas três superfícies sem mocks divergentes"),
            ("5 · Store Config / Onboarding", "Ativação da loja e criação progressiva de empresa", "Sequência cria company, membership, apps e dados iniciais isolados"),
            ("6 · Platform adapters", "Ports Web/Tauri e persistência", "Nenhum invoke/window.__TAURI__ espalhado em componentes"),
        ],
        [1600, 3510, 4250],
    )

    add_heading(doc, "12. Próximo vertical slice recomendado")
    add_status_callout(
        doc,
        "PRÓXIMO PASSO",
        "Implementar Catalog / Products de ponta a ponta, usando Galáxia Burger como fixture primária e Matriz Labs como tenant negativo de isolamento.",
        color=PURPLE_DARK,
        fill=LIGHT,
    )
    add_body(doc, "A sequência recomendada mantém arquitetura antes de tela: definir Product, ProductCategory e ProductModifier; escrever testes de leitura e mutação cruzada; criar repositório vinculado ao TenantContext; implementar calculateOrderItemPrice; produzir view models; então reproduzir a referência administrativa de Produtos e ligar disponibilidade ao futuro catálogo publicado.")
    add_bullets(doc, [
        "Teste primeiro: Galáxia Burger não lê nem altera produtos da Matriz Labs.",
        "Fonte única: o mesmo Product alimentará Admin, Store, Order e Dashboard.",
        "Pricing: preço base, modificadores e quantidade calculados no domínio/aplicação.",
        "UI: lista, categorias, busca, preço, estoque, disponibilidade, destaque e ações.",
        "Validação: testes, tipagem, lint, smoke global, execução e comparação visual desktop/mobile.",
    ])

    add_heading(doc, "13. Ponto de retomada")
    add_table(
        doc,
        ["Referência", "Valor"],
        [
            ("Branch integrada", "main"),
            ("Commit do merge", "90a377d · merge: tenant-safe Seumei Business OS hub"),
            ("Plano executado", "docs/superpowers/plans/2026-08-24-seumei-tenant-hub-foundation.md"),
            ("Especificação", "docs/superpowers/specs/2026-08-24-seumei-tenant-hub-foundation-design.md"),
            ("Conta demo", "demo@seumei.local"),
            ("Próxima branch sugerida", "codex/seumei-catalog-products"),
        ],
        [2500, 6860],
    )
    add_body(doc, "Este relatório é o checkpoint operacional: ele registra o que está publicado, quais garantias já existem, o que ainda não foi prometido como concluído e a ordem segura para continuar sem perder o contexto multiempresa.")

    core = doc.core_properties
    core.title = "Seumei Business OS — Progresso e Roadmap"
    core.subject = "Fundação multiempresa, shell autenticado e próximos vertical slices"
    core.author = "Matriz / Seumei"
    core.keywords = "Seumei, multitenancy, Business OS, Galáxia Burger, Matriz Labs"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
